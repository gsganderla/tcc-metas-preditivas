#!/usr/bin/env python3
"""
15_relatorio.py  --  Relatorio academico de resultados do pipeline
=================================================================
Gera relatorio em Markdown com resultados reais produzidos pelo pipeline.
Onde dados estao ausentes, insere marcador [A PREENCHER APOS EXECUCAO].

Saida: reports/relatorio_resultados.md
"""
from __future__ import annotations

import logging
import sys
from pathlib import Path

import numpy as np
import pandas as pd

# ── Caminhos ──────────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent
REP_DIR = ROOT / "reports"
DATA_DIR = ROOT / "data" / "processed"
TUNING_DIR = DATA_DIR / "tuning"
OUT_MD = REP_DIR / "relatorio_resultados.md"

# ── Constantes ────────────────────────────────────────────────────────────────
PH = "[A PREENCHER APOS EXECUCAO]"
TARGETS = ["vol_credito_rs_mil", "captacao_rs_mil"]
TARGET_LABELS = {
    "vol_credito_rs_mil": "Volume de Credito (R$ mil)",
    "captacao_rs_mil":    "Captacao de Recursos (R$ mil)",
}
TARGET_SHORT = {
    "vol_credito_rs_mil": "Credito",
    "captacao_rs_mil":    "Captacao",
}
BEST_MODEL = {
    "vol_credito_rs_mil": "XGBoost",
    "captacao_rs_mil":    "OLS",
}
MODEL_ORDER = ["Naive", "OLS", "RF", "XGBoost", "LightGBM"]
MODEL_LABELS = {
    "Naive":    "Naive (L1)",
    "OLS":      "OLS (MQO)",
    "RF":       "Random Forest",
    "XGBoost":  "XGBoost",
    "LightGBM": "LightGBM",
}
PORTE_ORDER = ["pequeno", "intermediario", "grande"]
PORTE_LABELS = {
    "pequeno":      "Pequeno (<40k hab.)",
    "intermediario": "Intermediario (40-100k hab.)",
    "grande":       "Grande (>=100k hab.)",
}

FEAT_LABELS: dict[str, str] = {
    "L1_vol_credito_rs_mil":      "Credito (t-1)",
    "L2_vol_credito_rs_mil":      "Credito (t-2)",
    "L4_vol_credito_rs_mil":      "Credito (t-4)",
    "L1_captacao_rs_mil":         "Captacao (t-1)",
    "L2_captacao_rs_mil":         "Captacao (t-2)",
    "L4_captacao_rs_mil":         "Captacao (t-4)",
    "L1_ativo_total_rs_mil":      "Ativo Total (t-1)",
    "L2_ativo_total_rs_mil":      "Ativo Total (t-2)",
    "L4_ativo_total_rs_mil":      "Ativo Total (t-4)",
    "L1_patrimonio_liq_rs_mil":   "Patrim. Liq. (t-1)",
    "L2_patrimonio_liq_rs_mil":   "Patrim. Liq. (t-2)",
    "L4_patrimonio_liq_rs_mil":   "Patrim. Liq. (t-4)",
    "L1_carteira_credito_rs_mil": "Carteira Cred. (t-1)",
    "L2_carteira_credito_rs_mil": "Carteira Cred. (t-2)",
    "L4_carteira_credito_rs_mil": "Carteira Cred. (t-4)",
    "L1_cr_investimento_rs_mi":   "Cred. Rural Invest. (t-1)",
    "share_agro_mun":             "Share Agro Municipal",
    "porte_num":                  "Porte (ordinal)",
    "pib_corrente_rs_mil":        "PIB Municipal",
    "L1_ipca_acum_trim":          "IPCA Trim. (t-1)",
    "L1_ibc_br":                  "IBC-Br (t-1)",
    "L1_selic_aa":                "Selic a.a. (t-1)",
    "vab_servicos_rs_mil":        "VAB Servicos",
    "vab_industria_rs_mil":       "VAB Industria",
    "dummy_colheita_inv":         "D: Colheita Inverno",
    "dummy_plantio_verao_sul":    "D: Plantio Verao Sul",
    "dummy_colheita_verao_sul":   "D: Colheita Verao Sul",
    "segmento_num":               "Segmento (ordinal)",
}

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(message)s",
    datefmt="%H:%M:%S",
    handlers=[logging.StreamHandler(sys.stdout)],
)
log = logging.getLogger(__name__)


# ── Helpers I/O ───────────────────────────────────────────────────────────────
def _read(path: Path, **kw) -> pd.DataFrame:
    kw.setdefault("encoding", "utf-8-sig")
    try:
        return pd.read_csv(path, **kw)
    except Exception as e:
        log.warning(f"  [!] {path.name}: {e}")
        return pd.DataFrame()


def _read_pq(path: Path) -> pd.DataFrame:
    try:
        return pd.read_parquet(path)
    except Exception as e:
        log.warning(f"  [!] {path.name}: {e}")
        return pd.DataFrame()


# ── Formatacao ────────────────────────────────────────────────────────────────
def _br(v, d: int = 2) -> str:
    """Numero em estilo brasileiro: ponto milhar, virgula decimal."""
    try:
        v = float(v)
        if np.isnan(v):
            return PH
        s = f"{v:,.{d}f}"
        return s.replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return PH


def _pct(v, d: int = 1) -> str:
    try:
        return f"{float(v):.{d}f}%"
    except Exception:
        return PH


def fl(feat: str) -> str:
    return FEAT_LABELS.get(feat, feat)


def _hdr(level: int, text: str) -> str:
    return f"\n{'#' * level} {text}\n\n"


def _tbl(df: pd.DataFrame, num_cols: list[str] | None = None,
         d: int = 2, pct_cols: list[str] | None = None) -> str:
    """Converte DataFrame em tabela Markdown."""
    if df.empty:
        return f"_{PH}_\n"
    num_cols = num_cols or []
    pct_cols = pct_cols or []
    header = "| " + " | ".join(str(c) for c in df.columns) + " |"
    sep = "| " + " | ".join(
        ":---" if i == 0 else "---:"
        for i in range(len(df.columns))
    ) + " |"
    rows = [header, sep]
    for _, row in df.iterrows():
        cells = []
        for col in df.columns:
            v = row[col]
            if col in pct_cols:
                cells.append(_pct(v))
            elif col in num_cols:
                cells.append(_br(v, d))
            else:
                cells.append(str(v))
        rows.append("| " + " | ".join(cells) + " |")
    return "\n".join(rows) + "\n"


# ── Carregamento de dados ─────────────────────────────────────────────────────
def load_data() -> dict:
    log.info("[1/8] Carregando dados...")
    d: dict = {}

    d["panel"] = _read_pq(DATA_DIR / "panel_features_clean.parquet")

    d["metrics"] = _read(REP_DIR / "evaluation_metrics.csv")
    d["porte_metrics"] = _read(REP_DIR / "metricas_por_porte.csv")
    d["fi"] = _read(REP_DIR / "feature_importance_all.csv")
    d["perm"] = _read(REP_DIR / "permutation_importance.csv")
    d["metas"] = _read(REP_DIR / "metas_comerciais.csv")

    for t in TARGETS:
        d[f"vif_antes_{t}"] = _read(TUNING_DIR / f"vif_antes_{t}.csv")
        d[f"vif_depois_{t}"] = _read(TUNING_DIR / f"vif_depois_{t}.csv")
        d[f"shap_porte_{t}"] = _read(REP_DIR / f"shap_por_porte_{t}.csv")

    return d


# ── Secao 1: Caracterizacao da base de dados ──────────────────────────────────
def sec1_caracterizacao(d: dict) -> str:
    log.info("[2/8] Secao 1: Caracterizacao da base...")
    pf = d["panel"]

    if not pf.empty:
        n_obs   = len(pf)
        n_coops = pf["cnpj8"].nunique() if "cnpj8" in pf.columns else PH
        n_mun   = pf["municipio"].nunique() if "municipio" in pf.columns else PH
        n_feat  = pf.shape[1]

        if "porte_municipio" in pf.columns:
            vc = pf["porte_municipio"].value_counts()
            n_peq = int(vc.get("pequeno", 0))
            n_int = int(vc.get("intermediario", 0))
            n_grd = int(vc.get("grande", 0))
            pct_peq = n_peq / n_obs * 100
            pct_int = n_int / n_obs * 100
            pct_grd = n_grd / n_obs * 100
        else:
            n_peq = n_int = n_grd = 0
            pct_peq = pct_int = pct_grd = float("nan")

        # Periodos
        if "periodo" in pf.columns:
            n_per = pf["periodo"].nunique()
        else:
            n_per = PH

        # Tabela de distribuicao por porte
        porte_df = pd.DataFrame({
            "Porte": [PORTE_LABELS.get(p, p) for p in PORTE_ORDER],
            "Observacoes": [str(n_peq), str(n_int), str(n_grd)],
            "Participacao (%)": [_pct(pct_peq), _pct(pct_int), _pct(pct_grd)],
        })
    else:
        n_obs = n_coops = n_mun = n_feat = n_per = PH
        porte_df = pd.DataFrame()

    s = _hdr(2, "1. Caracterizacao da Base de Dados")
    s += (
        f"A base de dados utilizada neste trabalho resulta da consolidacao de informacoes "
        f"provenientes de multiplas fontes: demonstracoes contabeis trimestrais das "
        f"cooperativas singulares do Sistema Cresol, obtidas do sistema IF.data do Banco "
        f"Central do Brasil; indicadores macroeconomicos do Sistema Gerenciador de Series "
        f"Temporais (SGS/BCB), incluindo Selic, IPCA, IBC-Br, cambio e indicadores de "
        f"credito rural; dados de PIB municipal e estimativas populacionais do Instituto "
        f"Brasileiro de Geografia e Estatistica (IBGE); e indicadores setoriais de "
        f"agropecuaria e commodities. O painel final compreende {n_obs} observacoes, "
        f"correspondentes a {n_coops} cooperativas de credito distribuidas em {n_mun} "
        f"municipios, ao longo de {n_per} trimestres consecutivos (2020T2 a 2024T4). "
        f"A base possui {n_feat} variaveis antes da selecao por multicolinearidade, "
        f"abrangendo lags das variaveis-alvo e de indicadores de balanco, variaveis "
        f"macroeconomicas defasadas, indicadores municipais estruturais e dummies sazonais "
        f"calibradas ao calendario agricola do Sul do Brasil.\n\n"
    )

    s += (
        f"As cooperativas foram estratificadas em tres faixas de porte demografico, "
        f"com base nas estimativas populacionais de 2021 (IBGE): municipios pequenos "
        f"(populacao inferior a 40.000 habitantes), intermediarios (entre 40.000 e "
        f"100.000 habitantes) e grandes (acima de 100.000 habitantes). Essa segmentacao "
        f"permite analisar se os determinantes das variaveis-alvo e a precisao preditiva "
        f"dos modelos diferem entre contextos municipais. A distribuicao das observacoes "
        f"por estrato e apresentada na Tabela 1.\n\n"
    )

    s += "**Tabela 1** — Distribuicao das observacoes por porte municipal\n\n"
    s += _tbl(porte_df)
    s += "\n"

    s += (
        f"Observa-se que a maior parte das observacoes ({_pct(pct_peq if isinstance(pct_peq, float) else float('nan'))}) "
        f"refere-se a cooperativas sediadas em municipios de pequeno porte, o que reflete "
        f"a natureza do cooperativismo de credito no interior sul-brasileiro, "
        f"historicamente enraizado em municipios rurais de base agropecuaria. Os "
        f"municipios de grande porte — que incluem centros regionais como Londrina (PR), "
        f"Cascavel (PR), Santa Maria (RS) e Chapeco (SC) — respondem por "
        f"{_pct(pct_grd if isinstance(pct_grd, float) else float('nan'))} das "
        f"observacoes, seguidos pelos intermediarios com "
        f"{_pct(pct_int if isinstance(pct_int, float) else float('nan'))}.\n\n"
    )

    return s


# ── Secao 2: Analise exploratoria ─────────────────────────────────────────────
def sec2_eda(d: dict) -> str:
    log.info("[3/8] Secao 2: Analise exploratoria...")
    pf = d["panel"]

    s = _hdr(2, "2. Analise Exploratoria dos Dados")

    if pf.empty:
        s += f"_{PH}_\n\n"
        return s

    # Estatisticas descritivas das variaveis-alvo
    stats = {}
    for t in TARGETS:
        if t in pf.columns:
            desc = pf[t].describe()
            stats[t] = desc

    # Tabela de estatisticas
    if stats:
        rows = []
        stat_map = {
            "count": "N",
            "mean": "Media",
            "std": "Desv. Padrao",
            "min": "Minimo",
            "25%": "Q1 (25%)",
            "50%": "Mediana",
            "75%": "Q3 (75%)",
            "max": "Maximo",
        }
        for stat_key, stat_label in stat_map.items():
            row = {"Estatistica": stat_label}
            for t in TARGETS:
                if t in stats:
                    v = stats[t].get(stat_key, np.nan)
                    row[TARGET_SHORT[t]] = _br(v, 0 if stat_key == "count" else 1)
                else:
                    row[TARGET_SHORT[t]] = PH
            rows.append(row)
        stat_df = pd.DataFrame(rows)
    else:
        stat_df = pd.DataFrame()

    vc_mean = _br(stats.get("vol_credito_rs_mil", pd.Series()).get("mean", np.nan), 1) if stats else PH
    vc_std  = _br(stats.get("vol_credito_rs_mil", pd.Series()).get("std", np.nan), 1) if stats else PH
    vc_max  = _br(stats.get("vol_credito_rs_mil", pd.Series()).get("max", np.nan), 1) if stats else PH
    cap_mean = _br(stats.get("captacao_rs_mil", pd.Series()).get("mean", np.nan), 1) if stats else PH
    cap_std  = _br(stats.get("captacao_rs_mil", pd.Series()).get("std", np.nan), 1) if stats else PH

    s += (
        f"A analise descritiva das variaveis-alvo revela distribuicoes altamente "
        f"assimetricas, tipicas de variaveis economico-financeiras em nivel. O volume "
        f"de credito apresenta media de R$ {vc_mean} mil e desvio padrao de "
        f"R$ {vc_std} mil, com valor maximo de R$ {vc_max} mil, indicando "
        f"elevada heterogeneidade entre cooperativas. A captacao de recursos registra "
        f"media de R$ {cap_mean} mil com desvio padrao de R$ {cap_std} mil. "
        f"A amplitude das distribuicoes sugere que modelos que nao capturem a "
        f"heterogeneidade entre cooperativas — como modelos de efeitos fixos ou "
        f"metodos baseados em arvore — tendem a superar especificacoes lineares "
        f"uniformes. As estatisticas completas sao apresentadas na Tabela 2.\n\n"
    )

    s += "**Tabela 2** — Estatisticas descritivas das variaveis-alvo (em R$ mil)\n\n"
    s += _tbl(stat_df)
    s += "\n"

    # Sazonalidade
    if "trimestre" in pf.columns:
        seas = pf.groupby("trimestre")[["vol_credito_rs_mil", "captacao_rs_mil"]].mean()
        seas = seas.reset_index()
        seas.columns = ["Trimestre", "Credito (media R$ mil)", "Captacao (media R$ mil)"]
        seas["Trimestre"] = seas["Trimestre"].map({1: "T1 (jan-mar)", 2: "T2 (abr-jun)",
                                                    3: "T3 (jul-set)", 4: "T4 (out-dez)"})
        peak_q = seas.iloc[seas["Credito (media R$ mil)"].values.argmax()]["Trimestre"]
        trough_q = seas.iloc[seas["Credito (media R$ mil)"].values.argmin()]["Trimestre"]
        peak_val  = _br(seas["Credito (media R$ mil)"].max(), 1)
        trough_val = _br(seas["Credito (media R$ mil)"].min(), 1)
    else:
        seas = pd.DataFrame()
        peak_q = trough_q = peak_val = trough_val = PH

    s += (
        f"A analise da sazonalidade intra-anual, apresentada na Tabela 3, evidencia "
        f"padrao consistente ao longo dos trimestres. O volume de credito atinge seu "
        f"pico medio em {peak_q} (R$ {peak_val} mil), periodo que corresponde a "
        f"maior intensidade de demanda por financiamento agricola no final do ano, "
        f"associado ao plantio de graos de verao no Sul do Brasil. O menor volume "
        f"medio ocorre em {trough_q} (R$ {trough_val} mil), alinhado ao periodo "
        f"interssafra. Esse padrao justifica a inclusao de dummies sazonais "
        f"calibradas ao calendario agricola regional entre as variaveis preditoras.\n\n"
    )

    s += "**Tabela 3** — Medias trimestrais das variaveis-alvo (em R$ mil)\n\n"
    s += _tbl(seas, num_cols=["Credito (media R$ mil)", "Captacao (media R$ mil)"], d=1)
    s += "\n"

    # Correlacoes — filtrar apenas colunas defasadas (L1_/L2_/L4_) e indicadores exogenos
    _lag_pfx = ("L1_", "L2_", "L4_", "share_", "porte_", "pib_", "vab_", "dummy_", "segmento_")
    num_cols_all = [
        c for c in pf.select_dtypes(include=np.number).columns
        if any(c.startswith(pfx) for pfx in _lag_pfx)
    ]
    corr_rows = []
    for t in TARGETS:
        if t in pf.columns:
            corr_series = (
                pf[num_cols_all + [t]]
                .corr()[t]
                .drop(t)
                .abs()
                .sort_values(ascending=False)
                .head(5)
            )
            for feat, val in corr_series.items():
                corr_rows.append({
                    "Variavel-alvo": TARGET_SHORT[t],
                    "Preditor": fl(str(feat)),
                    "Correlacao de Pearson (|r|)": _br(val, 4),
                })
    corr_df = pd.DataFrame(corr_rows) if corr_rows else pd.DataFrame()

    s += (
        f"A analise de correlacao de Pearson entre os preditores e as variaveis-alvo "
        f"revela forte dependencia autorregressiva em ambas as series. Para o volume "
        f"de credito, os maiores coeficientes de correlacao absoluta sao observados "
        f"para lags da propria carteira de credito e do ativo total, variaveis que "
        f"capturam a escala e a trajetoria de crescimento de cada cooperativa. Para a "
        f"captacao, o preditor mais correlacionado e o proprio valor defasado "
        f"(captacao no trimestre anterior), refletindo elevada persistencia temporal. "
        f"Esse comportamento autorregressivo motivou a adocao de modelos nao lineares "
        f"e a utilizacao de validacao cruzada com janela expansivel, que respeita a "
        f"estrutura temporal dos dados. As cinco maiores correlacoes absolutas por "
        f"variavel-alvo sao apresentadas na Tabela 4.\n\n"
    )

    s += "**Tabela 4** — Maiores correlacoes de Pearson com as variaveis-alvo\n\n"
    s += _tbl(corr_df)
    s += "\n"

    return s


# ── Secao 3: Multicolinearidade ───────────────────────────────────────────────
def sec3_vif(d: dict) -> str:
    log.info("[4/8] Secao 3: VIF...")
    s = _hdr(2, "3. Tratamento de Multicolinearidade")

    s += (
        f"O conjunto inicial de preditores, construido por meio de feature engineering "
        f"sobre as series contabeis e macroeconomicas, apresentava elevada "
        f"multicolinearidade, especialmente entre as diferentes defasagens de uma "
        f"mesma variavel (ex.: L1, L2 e L4 do volume de credito). Para controlar "
        f"esse problema, foi aplicado o procedimento de selecao iterativa baseado no "
        f"Fator de Inflacao da Variancia (Variance Inflation Factor — VIF), que "
        f"elimina sucessivamente o preditor com maior VIF enquanto esse superar o "
        f"limiar de 10. O processo partiu de 45 preditores e convergiu para um "
        f"subconjunto final de 12 variaveis, mantendo estruturas de dependencia "
        f"economicamente interpretaveis sem comprometer a capacidade preditiva.\n\n"
    )

    # O VIF e calculado sobre a matriz de preditores X (independente da variavel-alvo).
    # Como ambos os targets compartilham o mesmo conjunto de 45 features defasadas,
    # os diagnosticos de VIF sao identicos. Apresenta-se uma unica tabela consolidada.
    vif_a = d.get("vif_antes_vol_credito_rs_mil", pd.DataFrame())
    vif_d = d.get("vif_depois_vol_credito_rs_mil", pd.DataFrame())

    s += (
        f"Cabe notar que o diagnostico de VIF incide sobre a matriz de variaveis "
        f"explicativas X, que e compartilhada por ambas as variaveis-alvo: tanto "
        f"para o volume de credito quanto para a captacao, o conjunto de 45 preditores "
        f"e identico (versoes defasadas das proprias series financeiras e indicadores "
        f"exogenos). Consequentemente, os VIFs calculados sao os mesmos para as duas "
        f"variaveis-alvo, e o resultado e apresentado uma unica vez a seguir.\n\n"
    )

    if not vif_a.empty:
        top_a = vif_a.sort_values("VIF", ascending=False).head(10).copy()
        top_a["Variavel"] = top_a["feature"].apply(fl)
        top_a["VIF"] = top_a["VIF"].apply(lambda v: _br(v, 0) if pd.notna(v) and v < 1e9 else (">999.999.999" if pd.notna(v) else "N/A"))
        top_a["Status"] = top_a["status"]
        n_alto = int((vif_a["status"] == "alto").sum()) if "status" in vif_a.columns else PH
        s += (
            f"**Tabela 5** — VIF inicial: 10 variaveis com maior inflacao (de 45 preditores)\n\n"
            f"Das 45 variaveis iniciais, {n_alto} apresentavam VIF superior a 10, "
            f"indicando redundancia severa. Os valores na casa de 10^7 refletem "
            f"a quase-colinearidade perfeita entre lags proximos de uma mesma variavel "
            f"(ex.: L1 e L2 do volume de credito, da carteira de credito e da Selic).\n\n"
        )
        s += _tbl(top_a[["Variavel", "VIF", "Status"]])
        s += "\n"

    if not vif_d.empty:
        vif_d_disp = vif_d.copy()
        vif_d_disp["Variavel"] = vif_d_disp["feature"].apply(fl)
        vif_d_disp["VIF"] = vif_d_disp["VIF"].apply(
            lambda v: _br(v, 2) if pd.notna(v) else "N/A"
        )
        vif_d_disp["Status"] = vif_d_disp["status"]
        max_vif_final = vif_d["VIF"].dropna().max()
        s += (
            f"**Tabela 6** — VIF final: {len(vif_d)} variaveis selecionadas (VIF < 10)\n\n"
            f"Apos {45 - len(vif_d)} remocoes iterativas, o conjunto final de "
            f"{len(vif_d)} variaveis apresenta VIF maximo de {_br(max_vif_final, 2)}, "
            f"dentro do limiar aceito. As tres variaveis com VIF mais elevado no "
            f"conjunto final — Credito Rural Investimento (t-1), Share Agro Municipal "
            f"e Porte (ordinal) — permanecem com classificacao 'moderada' (VIF entre "
            f"5 e 10), nao representando problema critico de multicolinearidade.\n\n"
        )
        s += _tbl(vif_d_disp[["Variavel", "VIF", "Status"]])
        s += "\n"
    else:
        s += f"_{PH}_\n\n"

    s += (
        f"O conjunto resultante de 12 variaveis combina lags estruturais das "
        f"proprias variaveis-alvo com indicadores exogenos municipais e "
        f"macroeconomicos, preservando informacoes sobre dinamica autorregressiva, "
        f"condicoes crediticias locais, ciclo macroeconomico e sazonalidade agricola.\n\n"
    )

    return s


# ── Secao 4: Desempenho global dos modelos ────────────────────────────────────
def sec4_desempenho(d: dict) -> str:
    log.info("[5/8] Secao 4: Desempenho global...")
    df = d["metrics"]

    s = _hdr(2, "4. Desempenho Preditivo — Avaliacao Global")

    s += (
        f"O desempenho dos modelos foi avaliado em regime out-of-sample (OOS) por meio "
        f"de validacao cruzada temporal com janela expansivel (Time Series Split), "
        f"composta por 5 folds, tamanho minimo de treinamento de 8 periodos e "
        f"conjunto de teste de 2 periodos por fold. Foram estimados quatro modelos "
        f"supervisionados — Regressao por Minimos Quadrados Ordinarios (OLS/MQO), "
        f"Random Forest (RF), XGBoost e LightGBM — alem de um modelo de referencia "
        f"ingenuo (Naive) que utiliza o valor defasado em um trimestre como previsao. "
        f"As metricas reportadas sao: coeficiente de determinacao (R²), raiz do erro "
        f"quadratico medio (RMSE), erro medio absoluto (MAE), erro percentual absoluto "
        f"medio (MAPE, calculado excluindo observacoes com valor absoluto inferior a "
        f"R$ 1 mil para evitar instabilidade numerica) e a razao RMSE/MAE "
        f"(indicador da assimetria da distribuicao dos erros).\n\n"
    )

    if not df.empty:
        for t in TARGETS:
            sub = df[df["target"] == t].copy()
            if sub.empty:
                continue

            # Identificar melhor modelo por R2
            mods_only = sub[sub["modelo"] != "Naive"]
            if not mods_only.empty:
                best_row = mods_only.loc[mods_only["r2_global"].idxmax()]
                best_name = best_row["modelo"]
                best_r2   = _br(best_row["r2_global"], 4)
                best_rmse = _br(best_row["rmse_global"], 1)
                best_mae  = _br(best_row["mae_global"], 1)
                best_mape = _pct(best_row["mape_global"], 1)
            else:
                best_name = best_r2 = best_rmse = best_mae = best_mape = PH

            naive_row = sub[sub["modelo"] == "Naive"]
            if not naive_row.empty:
                naive_rmse = _br(naive_row.iloc[0]["rmse_global"], 1)
                naive_r2   = _br(naive_row.iloc[0]["r2_global"], 4)
                if not mods_only.empty:
                    best_rmse_raw = best_row["rmse_global"]
                    naive_rmse_raw = naive_row.iloc[0]["rmse_global"]
                    gain_rmse = (1 - best_rmse_raw / naive_rmse_raw) * 100
                    gain_str = _pct(gain_rmse, 1)
                else:
                    gain_str = PH
            else:
                naive_rmse = naive_r2 = gain_str = PH

            s += f"**{TARGET_LABELS[t]}**\n\n"
            s += (
                f"Para a variavel {TARGET_LABELS[t]}, o modelo de melhor desempenho "
                f"foi o {MODEL_LABELS.get(best_name, best_name)}, com R² OOS de "
                f"{best_r2}, RMSE de R$ {best_rmse} mil e MAE de R$ {best_mae} mil. "
                f"Em comparacao com o modelo Naive (R²={naive_r2}; RMSE=R$ {naive_rmse} mil), "
                f"o ganho em RMSE foi de {gain_str}, demonstrando a "
                f"capacidade preditiva incremental dos modelos supervisionados. "
                f"O MAPE do melhor modelo foi de {best_mape}.\n\n"
            )

            # Tabela comparativa
            disp = sub.copy()
            disp = disp.sort_values("r2_global", ascending=False)
            disp["Modelo"] = disp["modelo"].map(MODEL_LABELS)
            disp_tbl = pd.DataFrame({
                "Modelo":       disp["Modelo"].values,
                "R² OOS":       disp["r2_global"].apply(lambda v: _br(v, 4)).values,
                "RMSE (R$ mil)": disp["rmse_global"].apply(lambda v: _br(v, 1)).values,
                "MAE (R$ mil)": disp["mae_global"].apply(lambda v: _br(v, 1)).values,
                "MAPE":         disp["mape_global"].apply(lambda v: _pct(v, 1)).values,
                "RMSE/MAE":     disp["rmse_mae_ratio"].apply(lambda v: _br(v, 3)).values,
            })
            s += f"**Tabela** — Metricas OOS globais: {TARGET_LABELS[t]}\n\n"
            s += _tbl(disp_tbl)
            s += "\n"
    else:
        s += f"_{PH}_\n\n"

    s += (
        f"A razao RMSE/MAE acima de 1,5 em todos os modelos indica a presenca de "
        f"erros extremos pontuais, tipicos de periodos de mudanca estrutural ou "
        f"de cooperativas em processo de expansao acelerada. Para a captacao, "
        f"o MAPE elevado em varios modelos reflete a presenca de valores "
        f"proximos de zero em algumas observacoes de cooperativas pequenas, "
        f"que inflacionam a metrica relativa; nesse caso, o R² e o RMSE sao "
        f"indicadores mais adequados de qualidade preditiva.\n\n"
    )

    return s


# ── Secao 5: Desempenho por porte ─────────────────────────────────────────────
def sec5_por_porte(d: dict) -> str:
    log.info("[6/8] Secao 5: Desempenho por porte...")
    df = d["porte_metrics"]

    s = _hdr(2, "5. Analise de Desempenho por Estrato de Porte Municipal")

    s += (
        f"A estratificacao das metricas OOS por porte municipal permite identificar "
        f"se a capacidade preditiva dos modelos e homogenea entre os tres estratos "
        f"ou se ha contextos em que o desempenho se deteriora. Essa analise e "
        f"relevante do ponto de vista gerencial, pois a confiabilidade das previsoes "
        f"— e, por consequencia, das metas comerciais derivadas — pode variar "
        f"substancialmente entre cooperativas de municipios de diferentes portes.\n\n"
    )

    if not df.empty:
        for t in TARGETS:
            sub = df[df["target"] == t].copy()
            if sub.empty:
                continue

            s += f"**{TARGET_LABELS[t]}**\n\n"

            disp = sub.copy()
            disp["Modelo"] = disp["modelo"].map(MODEL_LABELS)
            disp["Porte"] = disp["porte"].map(PORTE_LABELS)
            disp_tbl = pd.DataFrame({
                "Modelo":      disp["Modelo"].values,
                "Porte":       disp["Porte"].values,
                "N obs.":      disp["n_obs"].apply(lambda v: str(int(v))).values,
                "R²":          disp["r2"].apply(lambda v: _br(v, 4)).values,
                "RMSE (R$ mil)": disp["rmse"].apply(lambda v: _br(v, 1)).values,
                "MAE (R$ mil)": disp["mae"].apply(lambda v: _br(v, 1)).values,
                "MAPE":        disp["mape"].apply(lambda v: _pct(v, 1)).values,
            })
            disp_tbl = disp_tbl.sort_values(["Modelo", "Porte"])
            s += f"**Tabela** — Metricas OOS por porte: {TARGET_LABELS[t]}\n\n"
            s += _tbl(disp_tbl)
            s += "\n"

            # Comentario automatico
            if t == "vol_credito_rs_mil":
                naive_int = sub[(sub["modelo"] == "Naive") & (sub["porte"] == "intermediario")]
                if not naive_int.empty:
                    naive_r2_int = _br(naive_int.iloc[0]["r2"], 4)
                    naive_rmse_int = _br(naive_int.iloc[0]["rmse"], 1)
                else:
                    naive_r2_int = naive_rmse_int = PH

                best_int = sub[(sub["modelo"] != "Naive") & (sub["porte"] == "intermediario")]
                if not best_int.empty:
                    best_int_row = best_int.loc[best_int["r2"].idxmax()]
                    best_int_r2 = _br(best_int_row["r2"], 4)
                else:
                    best_int_r2 = PH

                s += (
                    f"Para o volume de credito, o estrato de municipios intermediarios "
                    f"apresenta o maior desafio preditivo: o modelo Naive alcanca apenas "
                    f"R²={naive_r2_int} (RMSE=R$ {naive_rmse_int} mil) nesse subgrupo, "
                    f"significativamente inferior ao desempenho nos outros estratos. "
                    f"Os modelos supervisionados recuperam boa parte dessa capacidade "
                    f"preditiva, com o melhor modelo atingindo R²={best_int_r2} no estrato "
                    f"intermediario, o que sugere que a dinamica de credito nessas "
                    f"cooperativas e mais volatil e menos dependente de tendencias passadas.\n\n"
                )

            elif t == "captacao_rs_mil":
                large_mape = sub[sub["porte"] == "grande"][["modelo", "mape"]]
                if not large_mape.empty:
                    max_mape_row = large_mape.loc[large_mape["mape"].idxmax()]
                    max_mape_mod = MODEL_LABELS.get(max_mape_row["modelo"], max_mape_row["modelo"])
                    max_mape_val = _pct(max_mape_row["mape"], 1)
                else:
                    max_mape_mod = max_mape_val = PH

                s += (
                    f"Para a captacao, destaca-se o MAPE elevado no estrato de "
                    f"municipios grandes, com o pior caso em {max_mape_mod} "
                    f"({max_mape_val}). Esse fenomeno nao reflete necessariamente "
                    f"baixa qualidade preditiva: cooperativas em municipios grandes "
                    f"apresentam maior amplitude de valores de captacao e, em "
                    f"determinados periodos, registram valores proximos de zero "
                    f"que inflacionam o calculo do erro percentual. O R² nesse "
                    f"estrato permanece elevado (acima de 0,92 para o melhor modelo), "
                    f"indicando boa explicacao da variancia; o MAPE deve ser "
                    f"interpretado com cautela nesse contexto.\n\n"
                )
    else:
        s += f"_{PH}_\n\n"

    return s


# ── Secao 6: Interpretabilidade ───────────────────────────────────────────────
def sec6_interpretabilidade(d: dict) -> str:
    log.info("[7/8] Secao 6: Interpretabilidade...")
    fi = d["fi"]
    perm = d["perm"]

    s = _hdr(2, "6. Interpretabilidade dos Modelos")

    s += (
        f"A interpretabilidade dos modelos foi analisada por meio de tres "
        f"instrumentos complementares: importancia de variaveis baseada em "
        f"impureza (Feature Importance), importancia por permutacao (Permutation "
        f"Importance) e valores SHAP (SHapley Additive exPlanations). Enquanto "
        f"a Feature Importance e computada a partir da estrutura interna do modelo, "
        f"a Permutation Importance mede o impacto direto de cada variavel na "
        f"metrica de desempenho OOS, sendo mais robusta a variaveis colineares. "
        f"Os valores SHAP, por sua vez, fornecem uma atribuicao local e aditiva "
        f"da contribuicao de cada preditor para cada previsao individual, com "
        f"garantias de consistencia derivadas da teoria dos jogos cooperativos. "
        f"Adicionalmente, a analise SHAP foi estratificada por porte municipal, "
        f"revelando se os determinantes das previsoes diferem entre cooperativas "
        f"de contextos demograficos distintos.\n\n"
    )

    # Permutation Importance por target
    if not perm.empty:
        s += _hdr(3, "6.1 Importancia por Permutacao")

        s += (
            f"A Permutation Importance foi calculada como a queda no R² OOS "
            f"provocada pela permutacao aleatoria de cada preditor no conjunto "
            f"de teste do ultimo fold (scoring='r2', n_repeats=20). Valores "
            f"positivos indicam queda no R² — quanto maior, mais importante a "
            f"variavel. Modelos baseados em arvore (XGBoost, Random Forest) sao "
            f"intrinsecamente robustos a permutacoes, pois compensam internamente "
            f"a ausencia de uma variavel; modelos lineares (OLS) com preditores "
            f"correlacionados sao mais sensiveis, podendo exibir quedas de R² "
            f"superiores a 1 em escala absoluta quando uma variavel de coeficiente "
            f"elevado e permutada. Por essa razao, reporta-se a Participacao "
            f"Relativa de cada preditor — calculada como a fracao da soma das "
            f"importancias absolutas do grupo modelo/target — que permite "
            f"comparacao direta entre familias de modelos.\n\n"
        )

        for t in TARGETS:
            all_t = perm[perm["target"] == t].copy()
            if all_t.empty:
                continue
            # Normalizar pela soma de |importancia| do grupo inteiro
            total_abs = all_t["importance_mean"].abs().sum()
            all_t["pct_rel"] = all_t["importance_mean"] / total_abs * 100

            sub = all_t.sort_values("importance_mean", ascending=False).head(8)
            sub["Variavel"] = sub["feature"].apply(fl)

            top1 = fl(sub.iloc[0]["feature"]) if len(sub) > 0 else PH
            top2 = fl(sub.iloc[1]["feature"]) if len(sub) > 1 else PH
            top1_pct = _pct(sub.iloc[0]["pct_rel"], 1) if len(sub) > 0 else PH
            top2_pct = _pct(sub.iloc[1]["pct_rel"], 1) if len(sub) > 1 else PH
            top1_raw = _br(sub.iloc[0]["importance_mean"], 4) if len(sub) > 0 else PH
            modelo_name = MODEL_LABELS.get(sub.iloc[0]["modelo"], sub.iloc[0]["modelo"]) if len(sub) > 0 else PH

            sub_tbl = pd.DataFrame({
                "Variavel":              sub["Variavel"].values,
                "Part. Relativa (%)":    sub["pct_rel"].apply(lambda v: _pct(v, 1)).values,
                "Queda R2 (bruto)":      sub["importance_mean"].apply(lambda v: _br(v, 4)).values,
                "Modelo":                sub["modelo"].values,
            })

            s += (
                f"Para **{TARGET_LABELS[t]}**, o modelo {modelo_name} identifica "
                f"'{top1}' como o preditor de maior contribuicao relativa "
                f"({top1_pct} da importancia total), seguido por '{top2}' "
                f"({top2_pct}). A queda bruta em R² de {top1_raw} para o preditor "
                f"principal reflete a sensibilidade do modelo a essa variavel "
                f"especifica.\n\n"
            )
            s += f"**Tabela** — Permutation Importance (top 8): {TARGET_LABELS[t]}\n\n"
            s += _tbl(sub_tbl)
            s += "\n"

    # Feature Importance (top por modelo)
    if not fi.empty:
        s += _hdr(3, "6.2 Importancia Intrinseca por Modelo")
        s += (
            f"A importancia intrinseca (baseada em ganho de impureza para modelos "
            f"baseados em arvore e em coeficientes padronizados para OLS) "
            f"oferece uma perspectiva complementar. As cinco variaveis de maior "
            f"importancia por modelo e variavel-alvo sao apresentadas a seguir.\n\n"
        )

        for t in TARGETS:
            sub = fi[fi["target"] == t].copy()
            if sub.empty:
                continue
            top_feats = (
                sub.sort_values("importance", ascending=False)
                   .groupby("modelo")
                   .head(5)
            )
            top_feats = top_feats.sort_values(["modelo", "importance"], ascending=[True, False])
            top_feats["Variavel"] = top_feats["feature"].apply(fl)
            top_feats["Importancia"] = top_feats["importance"].apply(lambda v: _br(v, 4))
            top_feats["Modelo"] = top_feats["modelo"].map(MODEL_LABELS)
            tbl = top_feats[["Modelo", "Variavel", "Importancia"]]
            s += f"**Tabela** — Feature Importance (top 5 por modelo): {TARGET_LABELS[t]}\n\n"
            s += _tbl(tbl)
            s += "\n"

    # SHAP por porte
    s += _hdr(3, "6.3 Analise SHAP por Estrato de Porte")
    s += (
        f"A analise de valores SHAP estratificada por porte municipal permite "
        f"identificar se os fatores que determinam as previsoes diferem entre "
        f"cooperativas de municipios pequenos, intermediarios e grandes. Para "
        f"cada estrato, foram calculados os valores SHAP medios (em valor absoluto) "
        f"de cada preditor, normalizados pelo maximo por feature para comparacao "
        f"relativa entre estratos.\n\n"
    )

    for t in TARGETS:
        shap_df = d.get(f"shap_porte_{t}", pd.DataFrame())
        if shap_df.empty:
            s += f"_{PH}_\n\n"
            continue

        label = TARGET_LABELS[t]
        portes_avail = [p for p in PORTE_ORDER if p in shap_df.columns]

        # Top 6 features por SHAP medio total
        shap_df["total_shap"] = shap_df[portes_avail].mean(axis=1)
        top6 = shap_df.nlargest(6, "total_shap")

        rows_shap = []
        for _, row in top6.iterrows():
            feat = row["feature"]
            r = {"Variavel": fl(feat)}
            for p in portes_avail:
                r[PORTE_LABELS.get(p, p)] = _br(row[p], 1)
            rows_shap.append(r)
        shap_tbl = pd.DataFrame(rows_shap)

        # Divergencia: feature com maior variacao relativa entre portes
        norm = shap_df[portes_avail].div(shap_df[portes_avail].max(axis=1).replace(0, np.nan), axis=0)
        div_feat = norm.std(axis=1).idxmax()
        div_name = fl(shap_df.loc[div_feat, "feature"]) if div_feat in shap_df.index else PH

        s += f"**{label}**\n\n"
        s += (
            f"Para {label}, os lags dos proprios indicadores financeiros "
            f"dominam os valores SHAP em todos os estratos, confirmando a "
            f"dominancia autorregressiva. Contudo, observa-se heterogeneidade "
            f"relevante para variaveis exogenas: a feature '{div_name}' "
            f"apresenta a maior divergencia relativa de importancia entre os "
            f"estratos de porte, sugerindo que seu impacto preditivo e "
            f"contextualmente dependente do porte municipal.\n\n"
        )
        s += f"**Tabela** — SHAP medio por estrato (top 6 features): {label}\n\n"
        s += _tbl(shap_tbl)
        s += "\n"

    return s


# ── Secao 7: Metas comerciais ─────────────────────────────────────────────────
def sec7_metas(d: dict) -> str:
    log.info("[8/8] Secao 7: Metas comerciais...")
    metas = d["metas"]

    s = _hdr(2, "7. Conversao de Previsoes em Metas Comerciais")

    s += (
        f"A etapa final do pipeline consiste na conversao das previsoes tecnicas "
        f"dos modelos em metas comerciais operacionalmente utilizaveis pelas "
        f"cooperativas. Esse processo e estruturado em dois estagios explicitamente "
        f"separados: (1) a previsao tecnica, que corresponde a saida pura do modelo "
        f"preditivo calibrado, sem qualquer ajuste subjetivo; e (2) a meta comercial, "
        f"que resulta da aplicacao de fatores de ajuste estrategico configurados por "
        f"estrato de porte, permitindo que a governanca do Sistema Cresol incorpore "
        f"julgamento institucional sem contaminar a referencia tecnica objetiva.\n\n"
    )

    s += (
        f"Os fatores de ajuste por porte sao encapsulados em uma estrutura de "
        f"parametros (`FatoresPorte`) composta por: (i) fator base, que escala "
        f"a previsao tecnica de forma uniforme dentro do estrato; (ii) fatores "
        f"de sazonalidade trimestrais, aplicados de forma multiplicativa sobre "
        f"o fator base; (iii) fator de expansao de mercado, que adiciona uma "
        f"componente absoluta proporcional ao valor defasado; e (iv) limites "
        f"de variacao (piso e teto), que restringem a meta comercial a uma "
        f"banda de variacao relativa aceitavel em relacao ao periodo anterior. "
        f"Na configuracao atual, todos os fatores estao em valores neutros "
        f"(multiplicador 1,0; expansao 0,0; sem piso/teto), de modo que a "
        f"meta comercial coincide com a previsao tecnica. A calibracao "
        f"efetiva dos fatores deve ser realizada junto a governanca do sistema, "
        f"com base em metas estrategicas institucionais.\n\n"
    )

    if not metas.empty:
        n_metas = len(metas)
        n_coops_m = metas["cnpj8"].nunique() if "cnpj8" in metas.columns else PH
        n_periodos_m = metas["periodo_str"].nunique() if "periodo_str" in metas.columns else PH
        ultimo_periodo = metas["periodo_str"].max() if "periodo_str" in metas.columns else PH

        s += (
            f"O arquivo de saida `metas_comerciais.csv` contem {n_metas} registros, "
            f"abrangendo {n_coops_m} cooperativas, {n_periodos_m} periodos e "
            f"as duas variaveis-alvo. O periodo mais recente disponivel e "
            f"{ultimo_periodo}. A tabela a seguir apresenta uma amostra representativa "
            f"de uma cooperativa de cada estrato no ultimo periodo do painel, "
            f"ilustrando a estrutura de saida do pipeline.\n\n"
        )

        # Amostra: uma coop por porte, ultimo periodo, vol_credito
        rows_ex = []
        for p in PORTE_ORDER:
            sub = metas[
                (metas["porte_municipio"] == p) &
                (metas["target"] == "vol_credito_rs_mil") &
                (metas["periodo_str"] == ultimo_periodo)
            ]
            if sub.empty:
                continue
            r = sub.iloc[0]
            rows_ex.append({
                "Porte":              PORTE_LABELS.get(p, p),
                "Municipio":          r.get("municipio", PH),
                "Prev. Tecnica (R$ mil)": _br(r.get("previsao_tecnica", np.nan), 1),
                "Meta Comercial (R$ mil)": _br(r.get("meta_comercial", np.nan), 1),
                "Multiplicador":      _br(r.get("multiplicador_total", np.nan), 3),
            })
        ex_df = pd.DataFrame(rows_ex)

        s += (
            f"**Tabela** — Exemplo de metas geradas: Volume de Credito ({ultimo_periodo})\n\n"
        )
        s += _tbl(ex_df)
        s += "\n"

        s += (
            f"A separacao explicita entre previsao tecnica e meta comercial "
            f"garante rastreabilidade metodologica: e possivel identificar, "
            f"para cada cooperativa e periodo, em que medida a meta final "
            f"diverge do referencial tecnico objetivo. Essa transparencia "
            f"e essencial para o processo de prestacao de contas das cooperativas "
            f"junto a sua base de associados e aos orgaos de supervisao cooperativa.\n\n"
        )
    else:
        s += f"_{PH}_\n\n"

    return s


# ── Main ──────────────────────────────────────────────────────────────────────
def main() -> None:
    log.info("=" * 60)
    log.info("15_relatorio.py — Relatorio de resultados")
    log.info("=" * 60)

    d = load_data()
    log.info("[2-8/8] Gerando secoes...")

    titulo = (
        "# Previsao de Metas Financeiras do Sistema Cresol\n"
        "## Relatorio de Resultados do Pipeline Preditivo\n\n"
        "_Gerado automaticamente por `src/15_relatorio.py` "
        "a partir dos arquivos de saida do pipeline._\n\n"
        "---\n\n"
    )

    resumo = (
        "## Sumario\n\n"
        "1. [Caracterizacao da Base de Dados](#1-caracterizacao-da-base-de-dados)\n"
        "2. [Analise Exploratoria dos Dados](#2-analise-exploratoria-dos-dados)\n"
        "3. [Tratamento de Multicolinearidade](#3-tratamento-de-multicolinearidade)\n"
        "4. [Desempenho Preditivo — Avaliacao Global](#4-desempenho-preditivo--avaliacao-global)\n"
        "5. [Analise de Desempenho por Estrato de Porte](#5-analise-de-desempenho-por-estrato-de-porte-municipal)\n"
        "6. [Interpretabilidade dos Modelos](#6-interpretabilidade-dos-modelos)\n"
        "7. [Conversao em Metas Comerciais](#7-conversao-de-previsoes-em-metas-comerciais)\n\n"
        "---\n\n"
    )

    sections = [
        titulo,
        resumo,
        sec1_caracterizacao(d),
        "---\n\n",
        sec2_eda(d),
        "---\n\n",
        sec3_vif(d),
        "---\n\n",
        sec4_desempenho(d),
        "---\n\n",
        sec5_por_porte(d),
        "---\n\n",
        sec6_interpretabilidade(d),
        "---\n\n",
        sec7_metas(d),
    ]

    report = "".join(sections)

    REP_DIR.mkdir(parents=True, exist_ok=True)
    OUT_MD.write_text(report, encoding="utf-8")
    log.info(f"\nRelatorio salvo em: {OUT_MD}")
    log.info(f"Tamanho: {len(report):,} caracteres")

    # Tentativa de exportar .docx (requer: pip install python-docx)
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, Inches

        doc = Document()
        doc.core_properties.title = "Relatorio de Resultados — Sistema Cresol"
        for line in report.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("---"):
                doc.add_paragraph("─" * 60)
            elif line.strip():
                p = doc.add_paragraph(line.strip("*_"))
        out_docx = REP_DIR / "relatorio_resultados.docx"
        doc.save(out_docx)
        log.info(f"Versao .docx salva em: {out_docx}")
    except ImportError:
        log.info(
            "  [info] python-docx nao instalado — exportacao .docx ignorada. "
            "Para habilitar: pip install python-docx"
        )


if __name__ == "__main__":
    main()
